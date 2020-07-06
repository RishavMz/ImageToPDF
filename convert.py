#  Author : RishavMz
# A simple python program to create a docx file from images for further convertion into pdf format
# This program automises the boring procedure of copying all images and pasting in document which is very slow process,
#    and may also sometimes lead to freezing of your application. However, this program is extremely fast compared to 
#    performing the same task in MS Word. Moreover, this program resizes the images taking into account the aspect ratio
#    of the images which further minimizes the boring task of formatting each image one by one.    

#   Please read the readme file before executing the program 

from docx2pdf import convert
from docx import Document
from docx.shared import Inches, Cm
import cv2
import os
import inspect


document = Document()
size = list()
imagename = list()
path= os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
folder = path+'\images'

    # A loop to load imagename for all images which are to be inserted(in order)
for filename in os.listdir(folder):
    img = cv2.imread(os.path.join(folder,filename))
    if img is not None:
        print("Found ", filename)
        imagename.append(filename)
        image = img.shape
        width = img.shape[0]
        height = img.shape[1]
        size.append([width,height])

print('\n', len(imagename), 'images found')  
if(len(imagename) == 0):
    print("\n No image present in folder")
    quit(0)          
name = input("\nEnter a name for your pdf\n") 
name = name+'.docx'       

    # A loop to insert images from the folder into docx file one by one based on size and aspect ratio
for data in range(len(imagename)):
    print("Processing ",imagename[data])
    p = document.add_paragraph()
    r = p.add_run()
    #print(size[data])
    if (size[data][0] / size[data][1] >=1.5):
        #print("Type1")
        r.add_picture('images/'+(imagename[data]), height = Cm(27) , width = Cm(27/size[data][0]*size[data][1]))    
    elif (size[data][0] / size[data][1] >=1 ):
        #print("Type2")
        r.add_picture('images/'+(imagename[data]), height = Cm(24) , width = Cm(24/size[data][0]*size[data][1]))    
    else:
        #print("Type3")
        r.add_picture('images/'+(imagename[data]), width = Cm(21.4) , height = Cm(21.4/size[data][1]*size[data][0]))

widmargin = 0.1
lenmargin = 0.5    
sections = document.sections
    
    # Page margins narrowed
for section in sections:
    section.top_margin = Cm(lenmargin)
    section.bottom_margin = Cm(lenmargin)
    section.left_margin = Cm(widmargin)
    section.right_margin = Cm(widmargin)  
document.save(name)
print('\n==========\n=  Done  =\n==========\n')
print("Successfully created a docx file\n")
    
    # Converting the docx file to pdf file upon user's choice
print("\nPress 1 to convert to PDF and then press enter\n")
value = int(input())
if(value == 1): 
    print("\n Please wait until the process is completed....\nThis may take a minute\nYou will be displayed a success message after completion\n")
    convert(name,'PDF/',name[:-4]+".pdf")
no = input("\n\nSuccessfully converted to PDF format.\n\nPress enter key to continue")
