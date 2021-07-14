from docx import Document
from docx.shared import Inches, Cm
import cv2
import os, inspect


document = Document()
size = list()                   # Stores Image dimensions (for resizing into document)
imagename = list()              # Stores Images pathname which are to be added to docx file

path= os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))

imgpath = path+'\..\\temp\images'

for filename in os.listdir(imgpath):   # Storing data for indivisual images
    filepath = os.path.join(imgpath,filename)
    img = cv2.imread(filepath)
    if img is not None:
        print("Found ", filepath)
        imagename.append(filepath)
        image = img.shape
        width = img.shape[0]
        height = img.shape[1]
        size.append([width,height])

print('\n', len(imagename), 'images found')  
if(len(imagename) == 0):
    print("\n No image present in folder")
    quit(0)          
name = path + '\..\\temp\doc\sample.docx'      

for data in range(len(imagename)):      # Adding all images to docx file (along with proper resizing)
    print("Processing ",imagename[data])
    p = document.add_paragraph()
    r = p.add_run()
    if (size[data][1] / size[data][0] >= 1):
        r.add_picture(str(imagename[data]), height = Cm((20/size[data][1]) * size[data][0]) , width = Cm(20)  )  
    elif(size[data][0] / size[data][1] >= (25.94/20.59) ):
        r.add_picture(str(imagename[data]), height = Cm(25) , width = Cm(25/size[data][0]*size[data][1]))    
    else:
        r.add_picture(str(imagename[data]), width = Cm(20) , height = Cm(25))
widmargin = 0.5
lenmargintop = 1.5
lenmarginbottom = 0.5    
sections = document.sections
                   
for section in sections:                # Page margins narrowed
    section.top_margin = Cm(lenmargintop)
    section.bottom_margin = Cm(lenmarginbottom)
    section.left_margin = Cm(widmargin)
    section.right_margin = Cm(widmargin)  
document.save(name)
print('\n==========\n=  Done  =\n==========\n')
print("Successfully created a docx file\n")