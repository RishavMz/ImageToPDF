from docx import Document
from docx.shared import  Cm
from cv2 import imread
import os, inspect, sys
#import sqlite3

# Initializing variables
imageselected = list()
imagedimension = list()                   
imagepath = list()   
document = Document()
filename = sys.argv[1] 


# Initializing location and path 
path= os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
imageselectedpath = path+'\..\data\database\imageselected.txt'
dbpath = path+'\..\data\database\Image2PDF.db'
imgpath = path+'\..\data\images'
docxpath = path + '\..\data\doc\\'+filename+'.docx' 
   

fhand = open(imageselectedpath, "r")
for file in fhand:
    f = file.rstrip()
    if(len(f)>0):
        imageselected.append(f)
fhand.close()


## Add data into sqlite database (for future reference)
#dbconn = sqlite3.connect(dbpath)
#querydata = dbconn.execute("SELECT * FROM IMAGES;")
#for row in querydata:
#    imageselected.append(row[2])
#dbconn.close() 


# Collecting data about individual images
for filepath in imageselected:            
    img = imread(filepath)
    if img is not None:
        print("Found ", filepath)
        imagepath.append(filepath)
        image = img.shape
        width = img.shape[0]
        height = img.shape[1]
        imagedimension.append([width,height])


print('\n', len(imagepath), 'images found')  
if(len(imagepath) == 0):
    print("\n No image present in folder")
    quit(0)          

     
# Resizing images and adding them to docx file
for data in range(len(imagepath)):      
    print("Processing ",imagepath[data])
    p = document.add_paragraph()
    r = p.add_run()
    if (imagedimension[data][1] / imagedimension[data][0] >= 1):
        r.add_picture(str(imagepath[data]), height = Cm((20/imagedimension[data][1]) * imagedimension[data][0]) , width = Cm(20)  )  
    elif(imagedimension[data][0] / imagedimension[data][1] >= (25.94/20.59) ):
        r.add_picture(str(imagepath[data]), height = Cm(25) , width = Cm(25/imagedimension[data][0]*imagedimension[data][1]))    
    else:
        r.add_picture(str(imagepath[data]), width = Cm(20) , height = Cm(25))

                   
# Handling page margins for docx file
widmargin = 0.5
lenmargintop = 1.5
lenmarginbottom = 0.5    
sections = document.sections


for section in sections:                
    section.top_margin = Cm(lenmargintop)
    section.bottom_margin = Cm(lenmarginbottom)
    section.left_margin = Cm(widmargin)
    section.right_margin = Cm(widmargin)  
document.save(docxpath)


print('\n==========\n=  Done  =\n==========\n')
print("Successfully created a docx file\n")