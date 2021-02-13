#Author : RishavMz
#
#   This file provides an interactive GUI based application for the user to select images from his/her PC and 
#   convert it into docx format; and if required, then into PDF format.
#   
#   The GUI interface of this program is made using tkinker.
#   
#   Hope you like it :) ⭐🌟

import tkinter
from tkinter import *
from docx2pdf import convert
from docx import Document
from docx.shared import Inches, Cm
import cv2
import os
import inspect
from docx2pdf import convert
from tkinter.filedialog import askopenfilename

master = tkinter.Tk() 
master.title("Images to PDF converter")
state = 0

images_name = list()
def convertfn():
    global state
    document = Document()
    size = list()
    imagename = list()
    for filename in images_name:
        img = cv2.imread(os.path.join('',filename))
        if img is not None:
            print("Found ", filename)
            imagename.append(filename)
            image = img.shape
            width = img.shape[0]
            height = img.shape[1]
            size.append([width,height])
    print('\n', len(imagename), 'images found')  
    if(len(imagename) == 0):
        print("\n No image present in folder")
        label1 = Label(master,text = "No images selected!!! \nCreated a blank docx file\n")
        label1.pack()
    name = "sample"
    name = name+'.docx'
    print(name)
    for data in range(len(imagename)):
        print("Processing ",imagename[data])
        p = document.add_paragraph()
        r = p.add_run()
        #print(size[data])
        if (size[data][1] / size[data][0] >= 1):
        #    print("Type1 width >= height")
            r.add_picture(str(imagename[data]), height = Cm((20/size[data][1]) * size[data][0]) , width = Cm(20)  )  
        elif(size[data][0] / size[data][1] >= (25.94/20.59) ):
        #    print("Type2")
            r.add_picture(str(imagename[data]), height = Cm(25) , width = Cm(25/size[data][0]*size[data][1]))    
        else:
        #    print("Type3")
            r.add_picture(str(imagename[data]), width = Cm(20) , height = Cm(25))

    widmargin = 0.5
    lenmargintop = 1.5
    lenmarginbottom = 0.5    
    sections = document.sections
        
        # Page margins narrowed
    for section in sections:
        section.top_margin = Cm(lenmargintop)
        section.bottom_margin = Cm(lenmarginbottom)
        section.left_margin = Cm(widmargin)
        section.right_margin = Cm(widmargin)  
    document.save(name)
    print('\n==========\n=  Done  =\n==========\n')
    print("Successfully created a docx file\n")
    label1 = Label(master,text = "Successfully created a docx file\n")
    if(state == 1):  
        if(len(imagename)>0):
            label1.pack()
        state = 2
    
    


def createpdf1():
    global state
    print("Converted from docx to pdf format")

    
    print("\n Please wait until the process is completed....\nThis may take a minute\nYou will be displayed a success message after completion\n")
    convert("sample.docx",'PDF/',"sample.pdf")
    lab2 = Label(master , text = "Successfully converted PDF file.\n Check the PDF folder to find your file as sample.pdf \n")
    lab2.pack()
    print("\nSuccessfully created PDF file.\n Check the PDF folder to find your file as sample.pdf\n\n")

def createpdffn():
    global state
    lab1 = Label(master , text = " Please wait until the process is completed....\nThis may take a few seconds\nYou will be displayed a success message after completion\n\nNo need to click the continue button more than once.\n Sorry, but this process is a bit slow.\n")        
    button1 = tkinter.Button(master, text='Continue', width=15, command = createpdf1 )
    space = Label(text = "\n")
    if(state == 2):
        lab1.pack()
        button1.pack()
        space.pack()
        state = 3

def addimagesfn():

    global state
    print("Selecting images to add")
    while(True):
        path = askopenfilename() # show an "Open" dialog box and return the path to the selected file
        images_name.append(path)
        if(len(path)<=0):
            break
    if(len(images_name) == 0):
        print("\n No image present in folder")
        label1 = Label(adding,text = "No image selected\n")
    else:
        for data in range(len(images_name)):
            print("Addng", images_name[data])
        print('\n==========\n=  Done  =\n==========\n')
        label1 = Label(master,text = "Successfully added all images to memory\n")
    if(state == 0):    
        label1.pack()
        state = 1    



space = Label(text = "\n")
space.pack()
l1 = Label(master,text="This is a simple program to convert Images into docx and then into PDF file format.\n\nPlease select all the images you want to add to your PDF in your required order.\n\n")
l1.pack()
try:
    button1 = tkinter.Button(master, text='Add images to convert into PDF', width=35, command = addimagesfn)
    button1.pack()
    space = Label(text = "\n")
    space.pack()
    button1 = tkinter.Button(master, text='Convert Images to docx', width=35, command = convertfn)
    button1.pack()
    space = Label(text = "\n")
    space.pack()
    button1 = tkinter.Button(master, text='Convert docx file to PDF', width=35, command = createpdffn)
    button1.pack()
    space = Label(text = "\n")
    space.pack()
except:
    l2 = Label(master,text="An unexpected error has occured. Please contact us if this occurs repetedly\n")
    l2.pack()    

master.mainloop()
